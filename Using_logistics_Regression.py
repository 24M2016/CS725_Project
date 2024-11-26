import fitz  # PyMuPDF for PDF parsing
import os
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.svm import SVC
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import accuracy_score, classification_report
import camelot  # For table extraction
from PIL import Image
from io import BytesIO
import docx
from docx.shared import Inches
import matplotlib.pyplot as plt


# Function to extract text from a PDF using PyMuPDF
def extract_text_from_pdf(pdf_path):
    pdf_document = fitz.open(pdf_path)
    text = ""
    layout = []
    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)
        text += page.get_text("text")  # Extract text from page
        layout += page.get_text("dict")["blocks"]  # Extract layout blocks
    return text, layout


# Function to extract tables from a PDF
def extract_tables_from_pdf(pdf_path):
    try:
        tables = camelot.read_pdf(pdf_path, pages="all", flavor="stream")
        table_text = ""
        for table in tables:
            table_text += table.df.to_string(index=False) + "\n\n"
        return table_text
    except Exception as e:
        return ""  # Return empty string if no tables are found or an error occurs


# Function to extract images from a PDF
def extract_images_from_pdf(pdf_path):
    pdf_document = fitz.open(pdf_path)
    images = []
    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)
        for img_index, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            base_image = pdf_document.extract_image(xref)
            image_bytes = base_image["image"]
            images.append(Image.open(BytesIO(image_bytes)))
    return images


# Function to extract content from DOCX files (target data)
def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text


# Prepare dataset: Extract text, tables, images, and layout information
def prepare_data(pdf_dir, docx_dir):
    data = []
    for pdf_file in os.listdir(pdf_dir):
        if pdf_file.endswith(".pdf"):
            pdf_path = os.path.join(pdf_dir, pdf_file)
            docx_path = os.path.join(docx_dir, pdf_file.replace(".pdf", ".docx"))

            if os.path.exists(docx_path):
                # Extract text, tables, and images from PDF
                pdf_text, pdf_layout = extract_text_from_pdf(pdf_path)
                pdf_tables = extract_tables_from_pdf(pdf_path)
                pdf_images = extract_images_from_pdf(pdf_path)

                # Extract corresponding DOCX text
                docx_text = extract_text_from_docx(docx_path)

                # Append data: text + table + images + layout
                data.append({
                    "pdf_text": pdf_text,
                    "pdf_tables": pdf_tables,
                    "pdf_images": pdf_images,
                    "pdf_layout": pdf_layout,
                    "docx_text": docx_text,
                })
    return data


# Vectorize text content using TF-IDF
def vectorize_text(data):
    pdf_texts = [d["pdf_text"] + "\n" + d["pdf_tables"] for d in data]
    docx_texts = [d["docx_text"] for d in data]

    vectorizer = TfidfVectorizer(max_features=1000)
    X = vectorizer.fit_transform(pdf_texts).toarray()
    y = np.array([text.strip() for text in docx_texts])  # Flatten text to 1D array

    return X, y, vectorizer


# Train model to map PDF features to DOCX representation
def train_model(X, y, model_type="svm", epochs=10):
    accuracy = []
    precision = []
    recall = []
    f1 = []

    if model_type == "svm":
        model = SVC(kernel="linear")
    elif model_type == "logistic_regression":
        model = LogisticRegression(max_iter=1000)
    else:
        raise ValueError("Invalid model type. Choose 'svm' or 'logistic_regression'.")

    for epoch in range(epochs):
        model.fit(X, y)
        y_pred = model.predict(X)

        acc = accuracy_score(y, y_pred)
        accuracy.append(acc)

        report = classification_report(y, y_pred, output_dict=True)
        precision.append(report["weighted avg"]["precision"])
        recall.append(report["weighted avg"]["recall"])
        f1.append(report["weighted avg"]["f1-score"])

    return model, accuracy, precision, recall, f1


# Plot metrics for the model
def plot_metrics(epochs, accuracy, precision, recall, f1):
    plt.figure(figsize=(12, 8))
    plt.subplot(2, 2, 1)
    plt.plot(epochs, accuracy, marker="o", label="Accuracy", color="blue")
    plt.title("Accuracy")
    plt.grid(True)

    plt.subplot(2, 2, 2)
    plt.plot(epochs, precision, marker="o", label="Precision", color="green")
    plt.title("Precision")
    plt.grid(True)

    plt.subplot(2, 2, 3)
    plt.plot(epochs, recall, marker="o", label="Recall", color="purple")
    plt.title("Recall")
    plt.grid(True)

    plt.subplot(2, 2, 4)
    plt.plot(epochs, f1, marker="o", label="F1 Score", color="red")
    plt.title("F1 Score")
    plt.grid(True)

    plt.tight_layout()
    plt.show()


# Generate DOCX file from predictions
def create_docx_from_predictions(predictions, output_path):
    doc = docx.Document()
    doc.add_paragraph(predictions)
    doc.save(output_path)


# Main script
# Prepare data (Example paths)
docx_dir = "Dataset/docx"  # Folder containing PDF files
pdf_dir = "Dataset/pdfs"  # Folder containing corresponding DOCX files

data = prepare_data(pdf_dir, docx_dir)
X, y, vectorizer = vectorize_text(data)

model_type = "svm"
epochs = 10
model, accuracy, precision, recall, f1 = train_model(X, y, model_type, epochs)

plot_metrics(range(epochs), accuracy, precision, recall, f1)
