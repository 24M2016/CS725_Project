import os
import pandas as pd
from docx import Document
from pdf2docx import Converter
from PIL import Image
import pytesseract
import fitz  # PyMuPDF for PDFs
from tabula import read_pdf  # Requires tabula-py library


# Specify the Tesseract executable path (update as necessary)
pytesseract.pytesseract.tesseract_cmd = r'C:\Users\user\AppData\Local\Programs\Tesseract-OCR\tesseract.exe'


# Function to extract text from DOCX
def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    text = " ".join([para.text for para in doc.paragraphs])
    return text.strip()


# Function to extract tables from DOCX
def extract_tables_from_docx(docx_path):
    doc = Document(docx_path)
    tables_data = []

    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            tables_data.append(row_data)

    return tables_data


# Function to extract images from PDF
def extract_images_from_pdf(pdf_path, output_folder="extracted_images"):
    os.makedirs(output_folder, exist_ok=True)
    doc = fitz.open(pdf_path)
    images_info = []

    for i, page in enumerate(doc):
        for img_index, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            image_filename = f"{output_folder}/page_{i + 1}_img_{img_index + 1}.png"
            with open(image_filename, "wb") as img_file:
                img_file.write(image_bytes)
            images_info.append(image_filename)

    return images_info


# Function to extract text using OCR from PDF
def extract_text_with_ocr(pdf_path):
    doc = fitz.open(pdf_path)
    text = ""

    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        pix = page.get_pixmap()
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        text += pytesseract.image_to_string(img)

    return text.strip()


# Function to extract tables from PDF
def extract_tables_from_pdf(pdf_path):
    try:
        tables = read_pdf(pdf_path, pages="all", multiple_tables=True, lattice=True)
        return tables
    except Exception as e:
        return []


# Dataset Generation Function
def generate_dataset(pdf_folder, docx_folder, output_csv="dataset.csv"):
    data = []

    pdf_files = sorted([f for f in os.listdir(pdf_folder) if f.endswith(".pdf")])
    docx_files = sorted([f for f in os.listdir(docx_folder) if f.endswith(".docx")])

    if len(pdf_files) != len(docx_files):
        print("Warning: The number of PDF and DOCX files do not match!")

    for pdf_file, docx_file in zip(pdf_files, docx_files):
        pdf_path = os.path.join(pdf_folder, pdf_file)
        docx_path = os.path.join(docx_folder, docx_file)

        # Extract information from DOCX
        docx_text = extract_text_from_docx(docx_path)
        docx_tables = extract_tables_from_docx(docx_path)

        # Extract information from PDF
        pdf_text = extract_text_with_ocr(pdf_path)
        pdf_tables = extract_tables_from_pdf(pdf_path)
        pdf_images = extract_images_from_pdf(pdf_path)

        # Consolidate information
        data.append({
            "pdf_file": pdf_file,
            "docx_file": docx_file,
            "docx_text": docx_text,
            "pdf_text": pdf_text,
            "docx_tables": docx_tables,
            "pdf_tables": [table.to_string() for table in pdf_tables] if pdf_tables else None,
            "pdf_images": pdf_images
        })

    # Convert to DataFrame and Save as CSV
    df = pd.DataFrame(data)
    df.to_csv(output_csv, index=False)
    print(f"Dataset saved to {output_csv}")


# Example Usage
if __name__ == "__main__":
    pdf_folder_path = "D:/IITB Sem 1/CS 725/NewDataSet/TestPdfs"  # Replace with your PDF folder path
    docx_folder_path = "D:/IITB Sem 1/CS 725/NewDataSet/TestDocx"  # Replace with your DOCX folder path
    output_csv_path = "D:/IITB Sem 1/CS 725/NewDataSet/TEST_dataset.csv"

    generate_dataset(pdf_folder_path, docx_folder_path, output_csv_path)
